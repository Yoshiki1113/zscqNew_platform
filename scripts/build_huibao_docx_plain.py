# -*- coding: utf-8 -*-
"""生成《嘉剧荟》汇报材料通俗修订稿（少技术名词）。"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

OUT_DIR = Path(__file__).resolve().parents[1] / "docs" / "cailiao"
OUT_PATH = OUT_DIR / "“嘉剧荟”汇报材料（通俗易懂版）.docx"


def set_run_font(run, size=16, bold=False, name="仿宋"):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:eastAsia"), name)


def add_opening(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(0)
    pf.space_after = Pt(0)
    set_run_font(p.add_run(text), 16, False, "仿宋")


def add_heading_cn(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    set_run_font(p.add_run(text), 16, True, "黑体")


def add_body(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(0.74)
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    set_run_font(p.add_run(text), 16, False, "仿宋")


PARAS = [
    ("open", "各位领导："),
    ("body", "下面，我就“嘉剧荟”短剧侵权识别平台建设与应用情况作一简要汇报。"),
    (
        "body",
        "这套平台主要解决短剧版权保护中“发现难、取证难、协同难”的问题。简单说，就是让系统先帮我们找到疑似侵权内容、固定好证据材料，再由企业和公安按职责分工复核研判，形成一条清楚、可追踪的工作链条。",
    ),
    ("h", "一、项目建设背景"),
    (
        "body",
        "当前，短剧产业发展很快，传播也很快，已经成为文化产业的重要增长点。与此同时，侵权盗版问题日益突出。不少侵权内容出现在微信视频号等网络空间，更新快、删除也快，形式越来越多样，给版权保护和执法工作带来了新的压力。",
    ),
    ("body", "从实际工作看，主要有四个难点："),
    (
        "body",
        "一是发现难。平台内容海量，人工一条条翻找，既慢又容易漏，难以及时跟上内容更新节奏。",
    ),
    (
        "body",
        "二是认定难。现在不只是简单搬运，还有剪辑改标题、引流导流、换声改配等情况，单靠肉眼和经验判断，难度明显加大。",
    ),
    (
        "body",
        "三是取证难。很多内容“来得快、去得快”。特别是视频号场景，往往要进入播放页面，完成录屏、截图、账号信息留存，步骤多、窗口短，一旦中断，材料就不完整，后续核查和办案也会受影响。",
    ),
    (
        "body",
        "四是协同难。正版台词、侵权线索、取证材料、企业意见、公安研判信息分散在不同环节，缺少统一平台串起来，影响闭环处置。",
    ),
    (
        "body",
        "基于这些问题，我们建设“嘉剧荟”平台，就是想把大量重复、基础的工作交给系统，把人工精力集中到重点核查和最终研判上，真正提高工作效率和实战能力。",
    ),
    ("h", "二、平台总体定位"),
    (
        "body",
        "这套平台，可以理解为短剧版权保护的“智能助手”和“协同台账”。它不替代人工判断，而是先把活干在前面，把该留的证据留下来。",
    ),
    (
        "body",
        "对取证人员来说，平台可以按剧名去搜、去采，自动打开疑似侵权链接，完成录屏截图，并尽量把发布账号、引流去向等信息留住，减少反复手工操作。",
    ),
    (
        "body",
        "对企业来说，平台可以把证据材料整理好、比对意见标清楚，方便专人先复核，明确哪些更值得关注。",
    ),
    (
        "body",
        "对公安来说，平台把企业确认侵权的线索集中推送过来，形成可查看、可追踪的台账，减少来回沟通成本。",
    ),
    (
        "body",
        "一句话概括：系统先采集、先比对、先留痕；人再复核、再研判、再处置。",
    ),
    ("h", "三、怎么干：一条清晰工作链"),
    (
        "body",
        "我们按“谁发起、谁采集、谁复核、谁研判”来设计，整条链路可以概括为五步，环环相扣。",
    ),
    (
        "body",
        "第一步，企业报作品、交台词。公司把剧名清单和正版台词材料打包上传，形成工单。系统把台词整理进正版库，后面比对就有了“对照标准”。没有标准，就谈不上准确判断。",
    ),
    (
        "body",
        "第二步，把疑似侵权链接收进来。既可以按剧名自动搜索收集，也可以由企业把已经掌握的侵权链接直接提交。这些链接统一进链接池，便于排队处理、避免遗漏，也方便中途暂停后再接着做。",
    ),
    (
        "body",
        "第三步，对每条链接做完整取证。系统模拟手机真实操作，打开链接、进入播放页，完成录屏、截图，提取博主和引流相关信息。如果中途中断，支持接着采，并清理不完整的半截材料，保证最终留下来的是可用证据。",
    ),
    (
        "body",
        "第四步，听内容、对台词。系统把录屏里的声音转成文字，再和正版台词对照，给出相似程度和疑似等级。转写和比对也可以事后补做，不耽误现场先把证据固定下来。",
    ),
    (
        "body",
        "第五步，企业复核、公安研判。取证完成后，先推到公司核查池，由企业专人标注“侵权”或“未侵权”；确认侵权的，再推到公安端形成线索台账，做到发现一批、核查一批、沉淀一批。",
    ),
    ("h", "四、技术怎么支撑（少讲名词、多讲作用）"),
    (
        "body",
        "我们坚持“能落地、好维护、服务业务”的原则，不堆概念，重点看能不能把活干稳、材料留全。",
    ),
    (
        "body",
        "第一，有统一的业务平台。工单、任务、证据、复核、公安查看都在一个系统里完成，过程可看、结果可查，方便现场操作和汇报演示。",
    ),
    (
        "body",
        "第二，用手机自动化完成取证。视频号内容主要在手机里打开，系统就按真实操作路径去点、去录、去截，尽量把关键证据固定扎实，避免“看见了却没留下”。",
    ),
    (
        "body",
        "第三，用“听写+对台词”辅助判断。先把声音变成文字，再和正版台词比对，减少只靠标题猜的误判，让真正可疑的内容先浮出来。",
    ),
    (
        "body",
        "第四，自动提取页面关键信息。从截图里整理博主名称、账号标识、企业认证等信息，减少人工抄录，提高材料完整度。",
    ),
    (
        "body",
        "第五，材料和结论一起留存。系统不只给一个“像不像”的结论，还会同步保存录屏、截图、比对说明和报告，方便复查、会商和后续办案使用。",
    ),
    (
        "body",
        "总的看，这是一套“平台管流程、手机采证据、系统做比对、人工做决定”的组合拳，服务短剧版权保护实战。",
    ),
    ("h", "五、能够解决什么问题"),
    ("body", "从目前方向看，这套平台至少能在四个方面见效。"),
    (
        "body",
        "一是发现和采集更快。过去靠人工翻找、人工录屏，现在工单一牵、链接一进、系统跟采，能在更大范围、更短时间里形成可核查材料。",
    ),
    (
        "body",
        "二是研判更清楚。通过听写对台词、看引流去向、对照已知风险线索等方式综合判断，减少“标题像就算”的误报，让高风险线索优先出来。",
    ),
    (
        "body",
        "三是证据更完整。链接、截图、录屏、文字转写、比对结果都能留住，中断后还能续采，尽量避免材料残缺。",
    ),
    (
        "body",
        "四是协同更顺畅。企业发起、取证固化、企业复核、公安研判在同一链条上推进，台账清楚、责任清楚、结果可追踪。",
    ),
    ("h", "六、下一步工作方向"),
    ("body", "下一步，我们准备重点抓好五件事。"),
    (
        "body",
        "第一，继续充实正版台词和授权信息。对照标准越全越准，判断才会更扎实，并加强与辖区影视企业的对接。",
    ),
    (
        "body",
        "第二，稳妥扩大覆盖范围。当前以微信视频号为重点，后续视需要逐步延伸到更多传播渠道。",
    ),
    (
        "body",
        "第三，继续提高采集稳定性和识别能力。针对弹窗点击不稳、引流信息难抓、改头换面等难点，持续改进，努力少漏、少误。",
    ),
    (
        "body",
        "第四，继续打通实战闭环。让发起、采集、复核、推送、反馈各环节更顺畅，推动平台从“能用”走向“好用、管用”。",
    ),
    (
        "body",
        "第五，把流程和标准固化下来。逐步形成可复制的取证流程、报告模板和反馈机制，并视情完善权限管理等配套，便于推广应用。",
    ),
    ("h", "七、总结"),
    (
        "body",
        "总的来说，“嘉剧荟”不是单纯的技术展示，而是一套面向短剧版权保护实战的工作机制。",
    ),
    ("body", "它的核心价值，可以概括为三句话："),
    ("body", "第一，发现和取证更快；"),
    ("body", "第二，研判更清楚；"),
    ("body", "第三，协同闭环更有力。"),
    (
        "body",
        "下一步，我们将继续坚持问题导向、实战导向、应用导向，不断完善这套平台，努力把它建设成服务知识产权保护、支撑执法研判、助力企业维权的有效工具，为辖区文化产业健康发展提供更有力的保障。",
    ),
]


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    for kind, text in PARAS:
        if kind == "open":
            add_opening(doc, text)
        elif kind == "h":
            add_heading_cn(doc, text)
        else:
            add_body(doc, text)

    doc.save(str(OUT_PATH))
    full = "\n".join(t for _, t in PARAS)
    sys.stdout.reconfigure(encoding="utf-8")
    print(f"saved: {OUT_PATH}")
    print(f"chars: {len(full)}")


if __name__ == "__main__":
    main()
