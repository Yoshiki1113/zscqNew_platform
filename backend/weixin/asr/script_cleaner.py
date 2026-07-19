# -*- coding: utf-8 -*-
"""剧本台词清洗：doc/docx → 纯对白 txt（一行一句）

主路径：火山引擎 Ark 文本大模型；失败时规则兜底。
"""
from __future__ import annotations

import os
import re
import shutil
import subprocess
import tempfile
from io import BytesIO
from pathlib import Path
from typing import Optional

import httpx

try:
    from config import ARK_API_KEY, ARK_BASE_URL, ARK_MODEL, ARK_VISION_TIMEOUT
except ImportError:
    ARK_API_KEY = os.environ.get("ARK_API_KEY", "")
    ARK_BASE_URL = os.environ.get(
        "ARK_BASE_URL", "https://ark.cn-beijing.volces.com/api/plan/v3"
    )
    ARK_MODEL = os.environ.get("ARK_MODEL", "doubao-seed-2-0-mini-260428")
    ARK_VISION_TIMEOUT = int(os.environ.get("ARK_VISION_TIMEOUT", "60"))

# 文本清洗超时略长于视觉
ARK_SCRIPT_TIMEOUT = int(os.environ.get("ARK_SCRIPT_TIMEOUT", str(max(ARK_VISION_TIMEOUT, 120))))
ARK_SCRIPT_CHUNK_SIZE = int(os.environ.get("ARK_SCRIPT_CHUNK_SIZE", "4000"))
ARK_SCRIPT_CHUNK_OVERLAP = int(os.environ.get("ARK_SCRIPT_CHUNK_OVERLAP", "200"))
ARK_SCRIPT_MAX_TOKENS = int(os.environ.get("ARK_SCRIPT_MAX_TOKENS", "4096"))

_RE_SPEAKER = re.compile(r"^([^\s：:]{1,12})\s*[:：]\s*(.+)$")
_RE_EPISODE = re.compile(r"^第\d+集")
_RE_SCENE = re.compile(r"^\d+-\d+\s+")
_RE_CHARS = re.compile(r"^出场人物")
_RE_STAGE = re.compile(r"^▲")
_RE_SPEAKER_PREFIX = re.compile(r"^[\u4e00-\u9fffA-Za-z·]{1,12}\s*[:：]\s*")


SYSTEM_PROMPT = """你是短剧剧本清洗助手。任务：从剧本原文中只提取人物嘴里说的对白。

硬性规则：
1. 只输出对白正文，一句话一行。
2. 删除：角色名、`角色：`/`角色:` 前缀、集数/场次标题、出场人物、舞台指示、旁白说明、人物小传、场景描述。
3. 保留语气词与省略号（如 ……、咳咳……）。
4. 不要编号、不要 markdown、不要解释、不要空行。
5. 若本段没有对白，输出空（什么都不写）。
"""


def extract_docx_text(data: bytes) -> str:
    """从 docx 字节提取纯文本。"""
    from docx import Document

    doc = Document(BytesIO(data))
    parts: list[str] = []
    for para in doc.paragraphs:
        t = (para.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            line = "\t".join(c for c in cells if c)
            if line:
                parts.append(line)
    return "\n".join(parts)


def _convert_doc_via_win32com(doc_path: Path, out_docx: Path) -> bool:
    try:
        import pythoncom  # type: ignore
        import win32com.client  # type: ignore
    except ImportError:
        return False
    word = None
    com_inited = False
    try:
        # uvicorn/worker 线程下需手动初始化 COM
        pythoncom.CoInitialize()
        com_inited = True
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(str(doc_path.resolve()), ReadOnly=True)
        # 16 = wdFormatXMLDocument (.docx)
        doc.SaveAs2(str(out_docx.resolve()), FileFormat=16)
        doc.Close(False)
        return out_docx.is_file() and out_docx.stat().st_size > 0
    except Exception as e:
        print(f"[script_cleaner] win32com 转换失败: {e}")
        return False
    finally:
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        if com_inited:
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass


def _find_soffice() -> Optional[str]:
    candidates = [
        os.environ.get("PLATFORM_SOFFICE_PATH", "").strip(),
        shutil.which("soffice") or "",
        shutil.which("libreoffice") or "",
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]
    for c in candidates:
        if c and Path(c).exists():
            return c
    return None


def _convert_doc_via_libreoffice(doc_path: Path, out_dir: Path) -> Optional[Path]:
    soffice = _find_soffice()
    if not soffice:
        return None
    try:
        subprocess.run(
            [
                soffice,
                "--headless",
                "--convert-to",
                "docx",
                "--outdir",
                str(out_dir),
                str(doc_path),
            ],
            check=True,
            capture_output=True,
            timeout=120,
        )
        out = out_dir / (doc_path.stem + ".docx")
        if out.is_file() and out.stat().st_size > 0:
            return out
    except Exception as e:
        print(f"[script_cleaner] LibreOffice 转换失败: {e}")
    return None


def extract_doc_text(data: bytes) -> str:
    """从旧版 .doc（OLE）提取文本：先转 docx 再抽。"""
    with tempfile.TemporaryDirectory(prefix="script_doc_") as tmp:
        tmp_dir = Path(tmp)
        doc_path = tmp_dir / "input.doc"
        doc_path.write_bytes(data)
        out_docx = tmp_dir / "converted.docx"

        ok = _convert_doc_via_win32com(doc_path, out_docx)
        if not ok:
            converted = _convert_doc_via_libreoffice(doc_path, tmp_dir)
            if converted:
                out_docx = converted
                ok = True
        if not ok:
            raise RuntimeError(
                "无法解析 .doc：请安装 Microsoft Word 或 LibreOffice，或改为上传 .docx"
            )
        return extract_docx_text(out_docx.read_bytes())


def extract_script_bytes(data: bytes, filename: str = "") -> str:
    """按扩展名从剧本文件字节提取原文。"""
    name = (filename or "").lower()
    is_ole = len(data) >= 8 and data[:8] == b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"
    is_zip = len(data) >= 2 and data[:2] == b"PK"

    if name.endswith(".docx") or (is_zip and not name.endswith(".doc")):
        return extract_docx_text(data)
    if name.endswith(".doc") or is_ole:
        return extract_doc_text(data)
    if is_zip:
        return extract_docx_text(data)
    try:
        return extract_docx_text(data)
    except Exception:
        return extract_doc_text(data)


def _chunk_text(text: str, size: int = ARK_SCRIPT_CHUNK_SIZE, overlap: int = ARK_SCRIPT_CHUNK_OVERLAP) -> list[str]:
    text = text.replace("\r\n", "\n").replace("\r", "\n").strip()
    if not text:
        return []
    if len(text) <= size:
        return [text]

    chunks: list[str] = []
    start = 0
    n = len(text)
    while start < n:
        end = min(start + size, n)
        if end < n:
            # 优先在空行/换行处切开
            window = text[start:end]
            cut = max(window.rfind("\n\n"), window.rfind("\n"))
            if cut > size // 3:
                end = start + cut
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= n:
            break
        start = max(end - overlap, start + 1)
    return chunks


def _call_ark_text(prompt: str, user_content: str) -> Optional[str]:
    if not ARK_API_KEY:
        print("[script_cleaner] ARK_API_KEY 未配置")
        return None
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {ARK_API_KEY}",
    }
    payload = {
        "model": ARK_MODEL,
        "max_tokens": ARK_SCRIPT_MAX_TOKENS,
        "messages": [
            {"role": "system", "content": prompt},
            {"role": "user", "content": user_content},
        ],
        "thinking": {"type": "disabled"},
    }
    try:
        with httpx.Client(timeout=ARK_SCRIPT_TIMEOUT) as http:
            resp = http.post(
                f"{ARK_BASE_URL.rstrip('/')}/chat/completions",
                headers=headers,
                json=payload,
            )
        if resp.status_code != 200:
            print(f"[script_cleaner] HTTP {resp.status_code}: {resp.text[:300]}")
            return None
        result = resp.json()
        for choice in result.get("choices", []):
            content = choice.get("message", {}).get("content", "")
            if content:
                return content
        return None
    except Exception as e:
        print(f"[script_cleaner] API 调用失败: {e}")
        return None


def _rule_fallback_clean(raw: str) -> str:
    """规则兜底：尽量抽出「角色：台词」中的台词。"""
    lines_out: list[str] = []
    for raw_line in raw.replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        stripped = raw_line.strip()
        if not stripped:
            continue
        if _RE_EPISODE.match(stripped) or _RE_SCENE.match(stripped) or _RE_CHARS.match(stripped):
            continue
        if _RE_STAGE.match(stripped):
            continue
        m = _RE_SPEAKER.match(stripped)
        if m:
            dialog = m.group(2).strip()
            if dialog:
                lines_out.append(dialog)
            continue
        # 丢弃过长说明体
        if len(stripped) > 120 and "？" not in stripped and "！" not in stripped and "。" not in stripped[:40]:
            continue
        # 无冒号的短行也可能是对白
        if 2 <= len(stripped) <= 80:
            lines_out.append(stripped)
    return "\n".join(lines_out)


def _normalize_dialogue_lines(text: str) -> list[str]:
    lines: list[str] = []
    for raw_line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        s = raw_line.strip()
        if not s:
            continue
        # 去掉 markdown 噪音
        if s.startswith("```") or s in ("---", "***"):
            continue
        s = _RE_SPEAKER_PREFIX.sub("", s).strip()
        if not s:
            continue
        # 去掉行首序号
        s = re.sub(r"^\d+[\.、\)]\s*", "", s).strip()
        if s:
            lines.append(s)
    # 相邻完全重复去重
    deduped: list[str] = []
    for line in lines:
        if deduped and deduped[-1] == line:
            continue
        deduped.append(line)
    return deduped


def clean_script_to_dialogues(raw_text: str) -> tuple[str, str]:
    """清洗剧本原文为纯对白。

    Returns:
        (cleaned_text, mode)  mode 为 llm | fallback | mixed
    """
    raw_text = (raw_text or "").strip()
    if not raw_text:
        return "", "fallback"

    chunks = _chunk_text(raw_text)
    cleaned_parts: list[str] = []
    modes: set[str] = set()

    for i, chunk in enumerate(chunks):
        user_msg = (
            f"请清洗以下剧本片段（第 {i + 1}/{len(chunks)} 块），"
            f"只输出人物对白，一句话一行：\n\n{chunk}"
        )
        llm_out = _call_ark_text(SYSTEM_PROMPT, user_msg)
        if llm_out and llm_out.strip():
            cleaned_parts.append(llm_out)
            modes.add("llm")
        else:
            cleaned_parts.append(_rule_fallback_clean(chunk))
            modes.add("fallback")

    merged = "\n".join(cleaned_parts)
    lines = _normalize_dialogue_lines(merged)
    text = "\n".join(lines)
    if not text.strip():
        # 全失败再整篇规则兜底一次
        text = "\n".join(_normalize_dialogue_lines(_rule_fallback_clean(raw_text)))
        return text, "fallback"

    if modes == {"llm"}:
        mode = "llm"
    elif modes == {"fallback"}:
        mode = "fallback"
    else:
        mode = "mixed"
    return text, mode


def clean_script_file_bytes(data: bytes, filename: str = "") -> tuple[str, str, str]:
    """从文件字节抽原文并清洗。

    Returns:
        (cleaned_text, mode, raw_preview_len_info)
    """
    raw = extract_script_bytes(data, filename)
    cleaned, mode = clean_script_to_dialogues(raw)
    return cleaned, mode, f"raw_chars={len(raw)}"
